import csv
import re
import unicodedata
from pathlib import Path

from docx import Document

import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import os

import webbrowser

import sys
from pathlib import Path


BASE_DIR = Path(__file__).parent
def resource_path(relative_path):
    try:
        base_path = Path(sys._MEIPASS)
    except Exception:
        base_path = Path(__file__).parent

    return base_path / relative_path

CAMPOS = {
    "Formación y título profesional": {
        "required": True,
        "hidden": False,
        "type": "list"
    },
    "Área de trabajo": {
        "required": True,
        "hidden": False,
        "type": "text"
    },
    "Contacto": {
        "required": True,
        "hidden": False,
        "type": "email"
    },
    "Sede": {
        "required": True,
        "hidden": False,
        "type": "text"
    },
    "CV": {
        "required": False,
        "hidden": False,
        "type": "fixed_link"
    },
    "Trabajo final de doctorado": {
        "required": False,
        "hidden": True,
        "type": "text"
    },
    "Trabajo final de maestría": {
        "required": False,
        "hidden": True,
        "type": "text"
    },
    "Trabajo final de grado": {
        "required": False,
        "hidden": True,
        "type": "text"
    },
    "Scopus": {
        "required": False,
        "hidden": True,
        "type": "fixed_link"
    },
    "Orcid": {
        "required": False,
        "hidden": True,
        "type": "fixed_link"
    },
    "Linkedin": {
        "required": False,
        "hidden": True,
        "type": "fixed_link"
    }
}


def slugify(texto):
    texto = unicodedata.normalize("NFD", texto)
    texto = texto.encode("ascii", "ignore").decode("utf-8")
    texto = texto.lower()
    texto = re.sub(r"[^a-z0-9\s-]", "", texto)
    texto = re.sub(r"\s+", "-", texto.strip())
    return texto


def leer_docx(ruta):
    doc = Document(ruta)

    if not doc.tables:
        raise Exception("El documento no contiene tablas")

    tabla = doc.tables[0]

    datos = {}

    for fila in tabla.rows:

        if len(fila.cells) < 2:
            continue

        campo = fila.cells[0].text.strip()
        valor = fila.cells[1].text.strip()

        datos[campo] = valor

    return datos


def email_valido(email):
    patron = r"^[^@]+@[^@]+\.[^@]+$"
    return bool(re.match(patron, email))


def url_valida(url):
    return (
        url.startswith("http://")
        or
        url.startswith("https://")
    )


def generar_lista(valor):

    elementos = [
        linea.strip()
        for linea in valor.split("\n")
        if linea.strip()
    ]

    if not elementos:
        return ""

    html = "<ul>"

    for elemento in elementos:
        html += f"<li>{elemento}</li>"

    html += "</ul>"

    return html


def generar_celda(campo, valor, tipo):

    if tipo == "text":
        return valor

    if tipo == "list":
        return generar_lista(valor)

    if tipo == "email":
        return (
            f'<a href="mailto:{valor}" target="_blank">'
            f'{valor}'
            f'</a>'
        )

    if tipo == "fixed_link":
        return (
            f'<a href="{valor}" target="_blank">'
            f'{campo}'
            f'</a>'
        )

    return valor


def generar_fila(campo, valor, configuracion):

    atributos = ""

    if configuracion["hidden"]:
        atributos = ' class="hidden-row" style="display: none;"'

    contenido = generar_celda(
        campo,
        valor,
        configuracion["type"]
    )

    return f'''
<tr{atributos}>
    <td class="fichaColL">{campo}</td>
    <td class="fichaColR">{contenido}</td>
</tr>
'''


def generar_html(datos):

    filas = []
    hay_filas_ocultas = False

    for campo, configuracion in CAMPOS.items():

        valor = datos.get(campo, "").strip()

        if not configuracion["required"] and not valor:
            continue

        if configuracion["hidden"]:
            hay_filas_ocultas = True

        filas.append(
            generar_fila(
                campo,
                valor,
                configuracion
            )
        )

    html = '''
<table class="fichaIntegrantes">
    <tbody>
'''

    html += "".join(filas)

    html += '''
    </tbody>
</table>
'''

    if hay_filas_ocultas:

        html += '''
<button id="mostrarBotón" onclick="mostrarFilasOcultas()">Ver más</button>
<button id="ocultarBotón" onclick="ocultarFilas()" style="display:none;">Ver menos</button>
<script>
function mostrarFilasOcultas() {
    let hiddenRows = document.querySelectorAll('.hidden-row');
    hiddenRows.forEach(function(row) {
        row.style.display = 'table-row';
    });
    document.getElementById('mostrarBotón').style.display = 'none';
    document.getElementById('ocultarBotón').style.display = '';
}
function ocultarFilas() {
    let hiddenRows = document.querySelectorAll('.hidden-row');
    hiddenRows.forEach(function(row) {
        row.style.display = 'none';
    });
    document.getElementById('ocultarBotón').style.display = 'none';
    document.getElementById('mostrarBotón').style.display = '';
}
</script>
'''

    return html


def obtener_nombre_docente(datos):

    for clave in datos:

        if "Título abreviado" in clave:
            return datos[clave]

    return None


def procesar(archivos_seleccionados):

    primer_archivo = Path(archivos_seleccionados[0])

    carpeta_base = primer_archivo.parent

    SALIDA = carpeta_base / "Salida"

    REPORTE = carpeta_base / "reporte.csv"

    SALIDA.mkdir(exist_ok=True)

    registros = []

    cantidad_generados = 0

    for archivo in archivos_seleccionados:

        archivo = Path(archivo)

        try:

            datos = leer_docx(archivo)

            errores = []
            advertencias = []

            for campo, configuracion in CAMPOS.items():

                valor = datos.get(campo, "").strip()

                if configuracion["required"] and not valor:
                    errores.append(
                        f"Falta {campo}"
                    )

            email = datos.get(
                "Contacto",
                ""
            ).strip()

            if email and not email_valido(email):
                advertencias.append(
                    "Email inválido"
                )

            for campo in [
                "CV",
                "Scopus",
                "Orcid",
                "Linkedin"
            ]:

                valor = datos.get(
                    campo,
                    ""
                ).strip()

                if valor and not url_valida(valor):
                    advertencias.append(
                        f"URL inválida en {campo}"
                    )

            nombre_docente = obtener_nombre_docente(datos)

            if nombre_docente:

                nombre_archivo = (
                    slugify(nombre_docente)
                    + ".html"
                )

            else:

                nombre_archivo = (
                    slugify(archivo.stem)
                    + ".html"
                )

            html = generar_html(datos)

            ruta_salida = SALIDA / nombre_archivo

            ruta_salida.write_text(
                html,
                encoding="utf-8"
            )

            (SALIDA / nombre_archivo).write_text(
                html,
                encoding="utf-8"
            )

            cantidad_generados += 1

            estado = "OK"

            if errores:
                estado = "ERROR"

            elif advertencias:
                estado = "ADVERTENCIA"

            observaciones = []

            observaciones.extend(errores)
            observaciones.extend(advertencias)

            registros.append([
                archivo.name,
                estado,
                "; ".join(observaciones)
            ])

        except Exception as e:

            registros.append([
                archivo.name,
                "ERROR",
                str(e)
            ])

    with open(
        REPORTE,
        "w",
        newline="",
        encoding="utf-8"
    ) as archivo_csv:

        writer = csv.writer(
            archivo_csv
        )

        writer.writerow([
            "archivo",
            "estado",
            "observacion"
        ])

        writer.writerows(
            registros
        )

    return cantidad_generados, SALIDA

archivos_docx = []


archivos_docx = []


def seleccionar_archivos():

    global archivos_docx

    archivos_docx = filedialog.askopenfilenames(
        title="Seleccionar archivos DOCX",
        filetypes=[("Documentos Word", "*.docx")]
    )

    lista_archivos.delete(0, tk.END)

    for archivo in archivos_docx:
        lista_archivos.insert(tk.END, os.path.basename(archivo))


def ejecutar_proceso():

    if not archivos_docx:

        messagebox.showwarning(
            "Atención",
            "Debe seleccionar al menos un archivo DOCX."
        )

        return

    cantidad, carpeta_salida = procesar(
        archivos_docx
    )

    messagebox.showinfo(
        "Proceso finalizado",
        f"Se generaron {cantidad} fichas HTML.\n\n"
        f"Ubicación:\n{carpeta_salida}"
    )


    os.startfile(carpeta_salida)

def abrir_ayuda():
        webbrowser.open(
            "https://github.com/Gonzalosdesign/Generador-de-fichas-html?tab=readme-ov-file"
        )



if __name__ == "__main__":

    print("Creando ventana...")

    ventana = tk.Tk()

    ventana.iconphoto(
        True,
        tk.PhotoImage(
            file=resource_path("favicon.png")
        )
    )

    ventana.title(
        "Uneam Fagro UdelaR"
    )

    ventana.geometry(
        "600x400"
    )

    titulo = tk.Label(
        ventana,
        text="Generador de Fichas Docentes",
        font=("Gotham", 14, "bold"),
        fg="#33a033",
    )

    titulo.pack(
        pady=10
    )

    boton_seleccionar = tk.Button(
        ventana,
        text="Seleccionar archivos DOCX",
        command=seleccionar_archivos,
        font=("Gotham", 10),
        fg="white",
        bg="#304050"
    )

    boton_seleccionar.pack(
        pady=10
    )

    lista_archivos = tk.Listbox(
        ventana,
        width=80,
        height=12
    )

    lista_archivos.pack(
        padx=10,
        pady=10
    )

    boton_generar = tk.Button(
        ventana,
        text="Generar fichas HTML",
        command=ejecutar_proceso,
        bg="#4CAF50",
        fg="white",
        font=("Gotham", 10)
    )

    boton_generar.pack(
        pady=10
    )

    boton_generar.pack(
    pady=10
)

    boton_ayuda = tk.Button(
        ventana,
        text="ⓘ",
        command=abrir_ayuda,
        bg="#808080",
        fg="white",
        font=("Montserrat", 12, "bold"),

    )

    boton_ayuda.pack(
        side="bottom",
        anchor="se",
        padx=10,
        pady=10
    )

    ventana.mainloop()